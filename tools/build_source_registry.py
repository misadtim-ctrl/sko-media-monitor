#!/usr/bin/env python3
"""Build the normalized runtime source registry from the supplied DOCX."""

from __future__ import annotations

import argparse
import json
import re
import unicodedata
from collections import OrderedDict
from pathlib import Path
from urllib.parse import parse_qsl, urlencode, urlsplit, urlunsplit

from docx import Document

REPUBLICAN_TARGETS = [
    ("Tengrinews", "", "website", "https://tengrinews.kz/kazakhstan_news/", True, "лента новостей"),
    ("Zakon.kz", "", "website", "https://www.zakon.kz/", True, ""),
    ("Kazinform", "", "website", "https://www.inform.kz/ru", True, "русская версия"),
    ("Kazinform каз", "", "website", "https://www.inform.kz/kz", True, "казахская версия"),
    ("Nur.kz", "", "website", "https://www.nur.kz/", True, ""),
    ("Орда", "", "website", "https://orda.kz/", True, ""),
    ("Qumash.kz", "", "website", "https://qumash.kz/", True, ""),
    ("Hronika.kz", "", "website", "https://hronika.kz/", True, ""),
    ("Toppress.kz", "", "website", "https://toppress.kz/", True, ""),
    ("Best news", "", "website", "https://bestnews.kz/", True, "проверять диагностикой"),
    ("КазТаг", "", "website", "https://kaztag.kz/ru/", True, ""),
    ("Azattyq Ryh", "", "website", "https://rus.azattyq-ruhy.kz/", True, ""),
    ("Ratel", "", "website", "https://ratel.kz/", True, ""),
    ("Спутник", "", "website", "https://ru.sputnik.kz/", True, ""),
    ("BaigeNews.kz", "", "website", "https://baigenews.kz/", True, ""),
    (
        "Голос народа",
        "",
        "website",
        "https://golos-naroda.kz/",
        False,
        "адрес требует ручной проверки",
    ),
    ("EL.KZ", "", "website", "https://el.kz/ru/", True, ""),
    ("Inbusiness.kz", "", "website", "https://inbusiness.kz/ru", True, ""),
    ("Вести", "", "website", "https://vesti.kz/", True, ""),
    ("Тайм КЗ", "", "website", "https://www.time.kz/", True, ""),
    ("Тудей КЗ", "", "website", "https://today.kz/", False, "адрес требует ручной проверки"),
    ("Arna press", "", "website", "https://arnapress.kz/", True, ""),
    ("Қазақстан дәуірі", "", "website", "https://qazdauiri.kz/", True, ""),
    ("Егемен Қазақстан", "", "website", "https://egemen.kz/", True, ""),
    ("Казправда", "", "website", "https://www.kazpravda.kz/", True, ""),
    ("Литер", "", "website", "https://liter.kz/", True, ""),
    ("Exclusive", "", "website", "https://exclusive.kz/", True, ""),
    ("Baq.kz", "", "website", "https://rus.baq.kz/", True, ""),
    ("Ұлыс медиа", "", "website", "https://www.ulysmedia.kz/", True, ""),
    (
        "Халық үні",
        "",
        "website",
        "https://halyq-uni.kz/",
        False,
        "сайт был недоступен; оставить Telegram",
    ),
    ("Азаттық радиосы", "", "website", "https://rus.azattyq.org/", True, ""),
    ("NewTimes", "", "website", "https://newtimes.kz/", True, ""),
    ("Караван", "", "website", "https://www.caravan.kz/", True, ""),
    ("Kazlenta.kz", "", "website", "https://kazlenta.kz/feed/", True, "RSS"),
    ("Qaz365.kz", "", "website", "https://qaz365.kz/", True, ""),
    ("Первый канал Евразия", "", "youtube", "https://www.youtube.com/@1tveurasia", True, ""),
    ("31 канал", "", "youtube", "https://www.youtube.com/@31kanal", True, ""),
    ("Астана телеканал", "", "youtube", "https://www.youtube.com/@AstanaTV", True, ""),
    (
        "Qazaqstan TV",
        "",
        "youtube",
        "https://www.youtube.com/@qazaqstan_tv",
        True,
        "название Хабар в DOCX требует уточнения",
    ),
    ("24 KZ", "", "youtube", "https://www.youtube.com/@tv24kz", True, ""),
    ("Телеканал КТК", "", "youtube", "https://www.youtube.com/@ktk_kz", True, ""),
    ("Алматы ТВ", "", "youtube", "https://www.youtube.com/@almatytv", True, ""),
]

TRACKING_KEYS = {
    "fbclid",
    "yclid",
    "gclid",
    "igsh",
    "igshid",
    "mibextid",
    "ref",
    "ref_src",
    "from",
    "source",
    "share",
    "_r",
    "_t",
    "_rdr",
}
URL_RE = re.compile(r"https?://[^\s]+", re.IGNORECASE)


def normalize_url(raw: str) -> str:
    raw = raw.strip().rstrip(".,;)")
    if raw.startswith("https://web.telegram.org/") and "#@" in raw:
        raw = "https://t.me/" + raw.split("#@", 1)[1]
    parts = urlsplit(raw)
    host = parts.netloc.lower().removeprefix("www.")
    path = re.sub(r"/{2,}", "/", parts.path)
    if host == "t.me" and path.startswith("/s/"):
        path = path[2:]
    if host in {"instagram.com", "threads.com", "t.me", "vk.com", "tiktok.com"}:
        path = path.rstrip("/") + "/"
    query = [
        (key, value)
        for key, value in parse_qsl(parts.query, keep_blank_values=True)
        if not key.lower().startswith("utm_") and key.lower() not in TRACKING_KEYS
    ]
    canonical_host = {
        "instagram.com": "www.instagram.com",
        "youtube.com": "www.youtube.com",
        "tiktok.com": "www.tiktok.com",
    }.get(host, host)
    return urlunsplit(("https", canonical_host, path or "/", urlencode(query), ""))


def platform_for(url: str) -> str:
    host = urlsplit(url).netloc.lower().removeprefix("www.")
    if host == "t.me":
        return "telegram"
    if host == "instagram.com":
        return "instagram"
    if host == "threads.com":
        return "threads"
    if host == "vk.com":
        return "vk"
    if host == "facebook.com":
        return "facebook"
    if host == "tiktok.com":
        return "tiktok"
    if host == "youtube.com" or host == "youtu.be":
        return "youtube"
    return "website"


def slug(value: str) -> str:
    value = unicodedata.normalize("NFKD", value).lower()
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = re.sub(r"[^0-9a-zа-яё]+", "-", value, flags=re.IGNORECASE).strip("-")
    return value or "source"


def source_id(name: str, platform: str, url: str) -> str:
    path = urlsplit(url).path.strip("/").split("/")[0]
    return f"{platform}-{slug(path or name)}"


def add_source(registry: OrderedDict[str, dict], item: dict) -> None:
    key = item["url"].lower().rstrip("/")
    if key in registry:
        current = registry[key]
        current["aliases"] = sorted(set(current.get("aliases", []) + [item["name"]]))
        current["owners"] = sorted(set(current.get("owners", []) + item.get("owners", [])))
        current["docx_rows"] = sorted(set(current.get("docx_rows", []) + item.get("docx_rows", [])))
        return
    registry[key] = item


def build_registry(docx_path: Path) -> dict:
    doc = Document(docx_path)
    if len(doc.tables) < 2:
        raise ValueError("Expected two source tables in DOCX")

    registry: OrderedDict[str, dict] = OrderedDict()
    for name, _owner, platform, raw_url, enabled, notes in REPUBLICAN_TARGETS:
        url = normalize_url(raw_url)
        add_source(
            registry,
            {
                "id": source_id(name, platform, url),
                "name": name,
                "platform": platform,
                "url": url,
                "scope": "republican",
                "workflow": "sko_mentions",
                # Staff assignments from the internal DOCX are intentionally not
                # exported because the free GitHub runner uses a public repository.
                "owners": [],
                "enabled": enabled,
                "notes": notes,
                "aliases": [],
                "docx_rows": [],
            },
        )

    social_table = doc.tables[1]
    for row_number, row in enumerate(social_table.rows[1:], start=1):
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        name, _, links_cell, _owner = cells
        if row_number <= 28:
            scope, workflow = "local_public", "akimat_negative"
        elif row_number <= 35:
            scope, workflow = "regional", "regional_news"
        elif row_number <= 81:
            scope, workflow = "republican", "sko_mentions"
        else:
            scope, workflow = "civic_watch", "akimat_negative"

        for raw_url in URL_RE.findall(links_cell):
            url = normalize_url(raw_url)
            platform = platform_for(url)
            add_source(
                registry,
                {
                    "id": source_id(name, platform, url),
                    "name": name,
                    "platform": platform,
                    "url": url,
                    "scope": scope,
                    "workflow": workflow,
                    "owners": [],
                    "enabled": True,
                    "notes": "из рабочего списка РСК",
                    "aliases": [],
                    "docx_rows": [row_number],
                },
            )

    sources = list(registry.values())
    ids: dict[str, int] = {}
    for source in sources:
        base = source["id"]
        ids[base] = ids.get(base, 0) + 1
        if ids[base] > 1:
            source["id"] = f"{base}-{ids[base]}"

    counts: dict[str, int] = {}
    for source in sources:
        counts[source["platform"]] = counts.get(source["platform"], 0) + 1
    return {
        "schema_version": 1,
        "source_document": docx_path.name,
        "source_count": len(sources),
        "counts_by_platform": dict(sorted(counts.items())),
        "sources": sources,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    registry = build_registry(args.docx)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(registry, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(
        json.dumps(
            {
                "source_count": registry["source_count"],
                "counts_by_platform": registry["counts_by_platform"],
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
